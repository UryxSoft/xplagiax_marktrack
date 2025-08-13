from flask import Flask, request, jsonify, send_file, make_response,render_template
from flask_sqlalchemy import SQLAlchemy
from flask_caching import Cache
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from sqlalchemy import Column, Integer, Text, DateTime, String, ForeignKey
from minio import Minio
from minio.error import S3Error
import json
import sys
import gzip
import logging
import base64
import uuid
from datetime import datetime, timedelta
from io import BytesIO
import os
from urllib.parse import quote

# Librerías para exportación
from weasyprint import HTML, CSS
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from bs4 import BeautifulSoup

app = Flask(__name__)

# Configuración
class Config:
    SQLALCHEMY_DATABASE_URI = 'sqlite:///documents.db'
    SQLALCHEMY_TRACK_MODIFICATIONS = False
    MAX_DB_SIZE = 50000
    MAX_DOCUMENT_SIZE = 10 * 1024 * 1024  # 10MB
    AUTO_SAVE_DELAY = 2000
    KEEP_VERSIONS = 10
    CACHE_TYPE = 'simple'
    SECRET_KEY = 'tu-secret-key-aqui'

app.config.from_object(Config)

# Inicializar extensiones
db = SQLAlchemy(app)
cache = Cache(app)
#limiter = Limiter(key_func=get_remote_address)
#limiter.init_app(app)

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Cliente Minio
minio_client = Minio(
    'localhost:9500',
    access_key='minioadmin',
    secret_key='minioadmin',
    secure=False
)

# Crear buckets si no existen
BUCKETS = ['documents', 'images', 'exports', 'backups']
for bucket in BUCKETS:
    try:
        if not minio_client.bucket_exists(bucket):
            minio_client.make_bucket(bucket)
    except S3Error as e:
        logger.error(f"Error creando bucket {bucket}: {e}")

# Modelos de base de datos
class Document(db.Model):
    __tablename__ = 'documents'
    
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(255), nullable=False, default='Sin título')
    content_delta = db.Column(db.Text, nullable=True)
    content_html = db.Column(db.Text, nullable=True)
    minio_path = db.Column(db.String(255), nullable=True)
    storage_type = db.Column(db.String(20), default='database')
    size_bytes = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    versions = db.relationship('DocumentVersion', backref='document', cascade='all, delete-orphan')

class DocumentVersion(db.Model):
    __tablename__ = 'document_versions'
    
    id = db.Column(db.Integer, primary_key=True)
    document_id = db.Column(db.Integer, db.ForeignKey('documents.id'))
    version_number = db.Column(db.Integer)
    content_delta = db.Column(db.Text)
    minio_path = db.Column(db.String(255), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# Utilidades
def get_content_size(delta, html):
    """Calcula el tamaño del contenido"""
    delta_size = sys.getsizeof(json.dumps(delta))
    html_size = sys.getsizeof(html or '')
    return delta_size + html_size

def validate_delta(delta):
    """Valida que el delta sea válido"""
    if not isinstance(delta, dict):
        return False, "Delta debe ser un objeto JSON"
    
    if 'ops' not in delta:
        return False, "Delta debe contener 'ops'"
    
    # Verificar tamaño máximo
    if len(json.dumps(delta)) > Config.MAX_DOCUMENT_SIZE:
        return False, f"Documento excede tamaño máximo ({Config.MAX_DOCUMENT_SIZE} bytes)"
    
    return True, "Válido"

def extract_and_upload_images(delta):
    """Extrae imágenes del delta y las sube a Minio"""
    if not delta or 'ops' not in delta:
        return delta
    
    ops = delta.get('ops', [])
    
    for op in ops:
        if isinstance(op.get('insert'), dict):
            if 'image' in op['insert']:
                image_data = op['insert']['image']
                
                # Si es base64, subir a Minio
                if image_data.startswith('data:image'):
                    try:
                        # Extraer datos base64
                        header, base64_data = image_data.split(',', 1)
                        image_bytes = base64.b64decode(base64_data)
                        
                        # Generar nombre único
                        image_id = str(uuid.uuid4())
                        file_ext = 'png'  # Por defecto
                        if 'jpeg' in header:
                            file_ext = 'jpg'
                        elif 'gif' in header:
                            file_ext = 'gif'
                        
                        filename = f"{image_id}.{file_ext}"
                        
                        # Subir a Minio
                        minio_client.put_object(
                            bucket_name='images',
                            object_name=filename,
                            data=BytesIO(image_bytes),
                            length=len(image_bytes),
                            content_type=f'image/{file_ext}'
                        )
                        
                        # Reemplazar en delta con URL
                        op['insert']['image'] = f"/api/image/{filename}"
                        
                    except Exception as e:
                        logger.error(f"Error subiendo imagen: {e}")
    
    return delta

def save_to_minio_compressed(delta, html):
    """Guarda contenido comprimido en Minio"""
    content = {
        'delta': delta,
        'html': html,
        'timestamp': datetime.utcnow().isoformat()
    }
    
    json_content = json.dumps(content)
    compressed = gzip.compress(json_content.encode('utf-8'))
    
    filename = f"doc_{uuid.uuid4()}.json.gz"
    
    minio_client.put_object(
        bucket_name='documents',
        object_name=filename,
        data=BytesIO(compressed),
        length=len(compressed),
        content_type='application/gzip'
    )
    
    return filename

def load_from_minio_compressed(filename):
    """Carga contenido comprimido desde Minio"""
    try:
        response = minio_client.get_object('documents', filename)
        compressed_data = response.read()
        
        # Descomprimir
        json_content = gzip.decompress(compressed_data).decode('utf-8')
        content = json.loads(json_content)
        
        return content['delta'], content['html']
        
    except Exception as e:
        logger.error(f"Error cargando desde Minio: {e}")
        return None, None

def create_version(document):
    """Crea una nueva versión del documento"""
    # Mantener solo las últimas N versiones
    version_count = DocumentVersion.query.filter_by(document_id=document.id).count()
    
    if version_count >= Config.KEEP_VERSIONS:
        # Eliminar la versión más antigua
        oldest = DocumentVersion.query.filter_by(document_id=document.id)\
            .order_by(DocumentVersion.created_at)\
            .first()
        
        if oldest.minio_path:
            try:
                minio_client.remove_object('documents', oldest.minio_path)
            except:
                pass
        
        db.session.delete(oldest)
    
    # Crear nueva versión
    version = DocumentVersion(
        document_id=document.id,
        version_number=version_count + 1,
        content_delta=document.content_delta,
        minio_path=document.minio_path
    )
    
    db.session.add(version)

# Funciones de exportación
def clean_html_for_export(html):
    """Limpia HTML para exportación"""
    if not html:
        return ""
    
    soup = BeautifulSoup(html, 'html.parser')
    
    # Convertir imágenes de Minio a URLs completas
    for img in soup.find_all('img'):
        src = img.get('src', '')
        if src.startswith('/api/image/'):
            filename = src.replace('/api/image/', '')
            # En producción usarías la URL completa de Minio
            img['src'] = f"http://localhost:9000/images/{filename}"
    
    return str(soup)

def export_to_pdf(html_content, title="Documento"):
    """Exporta HTML a PDF usando WeasyPrint"""
    try:
        # CSS básico para el PDF
        css_content = """
        @page {
            margin: 2cm;
            @top-center {
                content: "{title}";
                font-family: Arial;
                font-size: 12px;
            }
        }
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            color: #333;
        }
        h1, h2, h3 { 
            color: #2c3e50; 
        }
        .ql-align-center { 
            text-align: center; 
        }
        .ql-align-right { 
            text-align: right; 
        }
        .ql-align-justify { 
            text-align: justify; 
        }
        blockquote {
            border-left: 4px solid #3498db;
            padding-left: 20px;
            margin-left: 0;
            font-style: italic;
        }
        """.format(title=title)
        
        html_doc = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{title}</title>
        </head>
        <body>
            {clean_html_for_export(html_content)}
        </body>
        </html>
        """
        
        pdf_buffer = BytesIO()
        HTML(string=html_doc).write_pdf(pdf_buffer, stylesheets=[CSS(string=css_content)])
        pdf_buffer.seek(0)
        
        return pdf_buffer
        
    except Exception as e:
        logger.error(f"Error exportando a PDF: {e}")
        return None

def export_to_docx(html_content, title="Documento"):
    """Exporta HTML a DOCX"""
    try:
        doc = DocxDocument()
        
        # Agregar título
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Procesar HTML
        soup = BeautifulSoup(clean_html_for_export(html_content), 'html.parser')
        
        # Función recursiva para procesar elementos
        def process_element(element, paragraph=None):
            if element.name == 'p':
                p = doc.add_paragraph()
                
                # Verificar alineación
                if 'ql-align-center' in element.get('class', []):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'ql-align-right' in element.get('class', []):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif 'ql-align-justify' in element.get('class', []):
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                process_children(element, p)
                
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                p = doc.add_heading(element.get_text(), level)
                
            elif element.name == 'blockquote':
                p = doc.add_paragraph(element.get_text())
                p.style = 'Quote'
                
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(), style='List Bullet')
                    
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(), style='List Number')
            
            elif element.name == 'br':
                if paragraph:
                    paragraph.add_run().add_break()
                else:
                    doc.add_paragraph()
        
        def process_children(element, paragraph):
            for child in element.children:
                if hasattr(child, 'name'):
                    if child.name == 'strong' or child.name == 'b':
                        run = paragraph.add_run(child.get_text())
                        run.bold = True
                    elif child.name == 'em' or child.name == 'i':
                        run = paragraph.add_run(child.get_text())
                        run.italic = True
                    elif child.name == 'u':
                        run = paragraph.add_run(child.get_text())
                        run.underline = True
                    elif child.name == 'br':
                        paragraph.add_run().add_break()
                    else:
                        paragraph.add_run(child.get_text())
                else:
                    paragraph.add_run(str(child))
        
        # Procesar todo el contenido
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'blockquote', 'ul', 'ol']):
            process_element(element)
        
        # Guardar en buffer
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        return docx_buffer
        
    except Exception as e:
        logger.error(f"Error exportando a DOCX: {e}")
        return None

# Rutas de la API
@app.route('/api/document', methods=['POST'])
#@limiter.limit("30/minute")
def create_document():
    """Crear nuevo documento"""
    data = request.get_json()
    
    title = data.get('title', 'Sin título')
    
    doc = Document(title=title)
    db.session.add(doc)
    db.session.commit()
    
    return jsonify({
        'id': doc.id,
        'title': doc.title,
        'created_at': doc.created_at.isoformat()
    })

@app.route('/api/document/<int:doc_id>/save', methods=['POST'])
#@limiter.limit("60/minute")
def save_document(doc_id):
    """Guardar documento con sistema híbrido"""
    data = request.get_json()
    
    if not data:
        return jsonify({'error': 'No se enviaron datos'}), 400
    
    delta = data.get('delta')
    html = data.get('html')
    
    if not delta:
        return jsonify({'error': 'Delta requerido'}), 400
    
    # Validar delta
    is_valid, message = validate_delta(delta)
    if not is_valid:
        return jsonify({'error': message}), 400
    
    # Buscar documento
    doc = Document.query.get_or_404(doc_id)
    
    # Crear versión antes de modificar
    create_version(doc)
    
    # Procesar imágenes
    delta = extract_and_upload_images(delta)
    
    # Calcular tamaño
    content_size = get_content_size(delta, html)
    doc.size_bytes = content_size
    doc.updated_at = datetime.utcnow()
    
    # Decidir almacenamiento
    if content_size <= Config.MAX_DB_SIZE:
        # Guardar en base de datos
        doc.content_delta = json.dumps(delta)
        doc.content_html = html
        doc.storage_type = 'database'
        
        # Limpiar Minio si existía
        if doc.minio_path:
            try:
                minio_client.remove_object('documents', doc.minio_path)
            except:
                pass
            doc.minio_path = None
    else:
        # Guardar en Minio
        minio_path = save_to_minio_compressed(delta, html)
        doc.minio_path = minio_path
        doc.storage_type = 'minio'
        doc.content_delta = None
        doc.content_html = None
    
    db.session.commit()
    
    # Limpiar caché
    cache.delete(f'document_{doc_id}')
    
    logger.info(f"Documento {doc_id} guardado en {doc.storage_type}, tamaño: {content_size}")
    
    return jsonify({
        'status': 'saved',
        'storage_type': doc.storage_type,
        'size_bytes': content_size,
        'updated_at': doc.updated_at.isoformat()
    })

@app.route('/api/document/<int:doc_id>/load', methods=['GET'])
@cache.memoize(timeout=300)
def load_document(doc_id):
    """Cargar documento"""
    doc = Document.query.get_or_404(doc_id)
    
    if doc.storage_type == 'database':
        delta = json.loads(doc.content_delta) if doc.content_delta else {}
        html = doc.content_html or ''
    else:
        delta, html = load_from_minio_compressed(doc.minio_path)
        if delta is None:
            return jsonify({'error': 'Error cargando desde almacenamiento'}), 500
    
    return jsonify({
        'id': doc.id,
        'title': doc.title,
        'delta': delta,
        'html': html,
        'storage_type': doc.storage_type,
        'size_bytes': doc.size_bytes,
        'created_at': doc.created_at.isoformat(),
        'updated_at': doc.updated_at.isoformat()
    })

@app.route('/api/document/<int:doc_id>/export/<format_type>', methods=['GET'])
def export_document(doc_id, format_type):
    """Exportar documento a PDF o DOCX"""
    if format_type not in ['pdf', 'docx']:
        return jsonify({'error': 'Formato no soportado'}), 400
    
    doc = Document.query.get_or_404(doc_id)
    
    # Cargar contenido
    if doc.storage_type == 'database':
        html = doc.content_html or ''
    else:
        delta, html = load_from_minio_compressed(doc.minio_path)
        if html is None:
            return jsonify({'error': 'Error cargando documento'}), 500
    
    # Generar archivo
    if format_type == 'pdf':
        file_buffer = export_to_pdf(html, doc.title)
        mimetype = 'application/pdf'
        filename = f"{doc.title}.pdf"
    else:  # docx
        file_buffer = export_to_docx(html, doc.title)
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        filename = f"{doc.title}.docx"
    
    if not file_buffer:
        return jsonify({'error': 'Error generando archivo'}), 500
    
    try:
        # Guardar en Minio para descarga
        export_filename = f"{doc_id}_{uuid.uuid4().hex}.{format_type}"
        
        minio_client.put_object(
            bucket_name='exports',
            object_name=export_filename,
            data=file_buffer,
            length=file_buffer.getbuffer().nbytes,
            content_type=mimetype
        )
        
        # Generar URL de descarga temporal
        download_url = minio_client.presigned_get_object(
            'exports', 
            export_filename, 
            expires=timedelta(hours=1)
        )
        
        return jsonify({
            'download_url': download_url,
            'filename': filename,
            'expires_in': '1 hora'
        })
        
    except Exception as e:
        logger.error(f"Error exportando documento: {e}")
        return jsonify({'error': 'Error generando exportación'}), 500

@app.route('/api/image/<filename>')
def serve_image(filename):
    """Servir imágenes desde Minio"""
    try:
        response = minio_client.get_object('images', filename)
        
        # Determinar tipo de contenido
        content_type = 'image/png'  # por defecto
        if filename.lower().endswith('.jpg') or filename.lower().endswith('.jpeg'):
            content_type = 'image/jpeg'
        elif filename.lower().endswith('.gif'):
            content_type = 'image/gif'
        
        return send_file(
            BytesIO(response.read()),
            mimetype=content_type,
            as_attachment=False
        )
        
    except Exception as e:
        logger.error(f"Error sirviendo imagen {filename}: {e}")
        return jsonify({'error': 'Imagen no encontrada'}), 404

@app.route('/api/documents', methods=['GET'])
def list_documents():
    """Listar todos los documentos"""
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    
    docs = Document.query.order_by(Document.updated_at.desc())\
        .paginate(page=page, per_page=per_page, error_out=False)
    
    return jsonify({
        'documents': [{
            'id': doc.id,
            'title': doc.title,
            'storage_type': doc.storage_type,
            'size_bytes': doc.size_bytes,
            'created_at': doc.created_at.isoformat(),
            'updated_at': doc.updated_at.isoformat()
        } for doc in docs.items],
        'total': docs.total,
        'pages': docs.pages,
        'current_page': page
    })

@app.route('/api/stats', methods=['GET'])
def get_stats():
    """Estadísticas del sistema"""
    total_docs = Document.query.count()
    db_docs = Document.query.filter_by(storage_type='database').count()
    minio_docs = Document.query.filter_by(storage_type='minio').count()
    
    total_size = db.session.query(db.func.sum(Document.size_bytes)).scalar() or 0
    
    return jsonify({
        'total_documents': total_docs,
        'database_documents': db_docs,
        'minio_documents': minio_docs,
        'total_size_bytes': total_size,
        'total_size_mb': round(total_size / (1024 * 1024), 2)
    })


@app.route('/')
def index():
    """Route principal que renderiza el template index.html"""
    return render_template('frontend.html')
# Inicialización
@app.before_request
def initialize_database():
    if not hasattr(app, '_database_initialized'):
        db.create_all()
        app._database_initialized = True
if __name__ == '__main__':
    
    app.run(debug=True,host='127.0.0.1', port=5002)
