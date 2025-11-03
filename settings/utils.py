import json
import sys
import gzip
import uuid
import base64
import logging
import re
from datetime import datetime, timedelta
from io import BytesIO
from urllib.parse import quote
from werkzeug.utils import secure_filename

# Librerías para procesamiento de documentos
import mammoth
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML, CSS
from bs4 import BeautifulSoup, Tag, NavigableString
from PIL import Image
from models.models import DocumentVersion

# Flask y extensiones
from flask import current_app, request
from flask_mail import Message

from .extensions import minio_client, redis_client, mail, logger,db

def allowed_file(filename, allowed_extensions=None):
    """Verificar si el archivo tiene una extensión permitida"""
    if allowed_extensions is None:
        allowed_extensions = current_app.config['ALLOWED_EXTENSIONS']
    
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions

def generate_safe_filename(filename):
    """Generar nombre de archivo seguro"""
    filename = secure_filename(filename)
    name, ext = filename.rsplit('.', 1) if '.' in filename else (filename, '')
    unique_id = str(uuid.uuid4())[:8]
    return f"{name}_{unique_id}.{ext}" if ext else f"{name}_{unique_id}"

def get_content_size(delta, html):
    """Calcular el tamaño del contenido"""
    delta_size = sys.getsizeof(json.dumps(delta))
    html_size = sys.getsizeof(html or '')
    return delta_size + html_size

def validate_delta(delta):
    """Validar que el delta sea válido"""
    if not isinstance(delta, dict):
        return False, "Delta debe ser un objeto JSON"
    
    if 'ops' not in delta:
        return False, "Delta debe contener 'ops'"
    
    # Verificar tamaño máximo
    if len(json.dumps(delta)) > current_app.config['MAX_DOCUMENT_SIZE']:
        return False, f"Documento excede tamaño máximo ({current_app.config['MAX_DOCUMENT_SIZE']} bytes)"
    
    return True, "Válido"

def validate_email(email):
    """Validar formato de email"""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

def extract_and_upload_images(delta):
    """Extraer imágenes del delta y subirlas a Minio"""
    if not delta or 'ops' not in delta:
        return delta
    
    ops = delta.get('ops', [])
    
    for op in ops:
        if isinstance(op.get('insert'), dict):
            if 'image' in op['insert']:
                image_data = op['insert']['image']
                
                # Si es base64, subir a Minio
                if isinstance(image_data, str) and image_data.startswith('data:image'):
                    try:
                        # Extraer datos base64
                        header, base64_data = image_data.split(',', 1)
                        image_bytes = base64.b64decode(base64_data)
                        
                        # Generar nombre único
                        image_id = str(uuid.uuid4())
                        file_ext = 'png'  # Por defecto
                        if 'jpeg' in header or 'jpg' in header:
                            file_ext = 'jpg'
                        elif 'gif' in header:
                            file_ext = 'gif'
                        elif 'webp' in header:
                            file_ext = 'webp'
                        
                        filename = f"{image_id}.{file_ext}"
                        
                        # Optimizar imagen si es necesario
                        optimized_bytes = optimize_image(image_bytes, file_ext)
                        
                        # Subir a Minio
                        minio_client.put_object(
                            bucket_name='images',
                            object_name=filename,
                            data=BytesIO(optimized_bytes),
                            length=len(optimized_bytes),
                            content_type=f'image/{file_ext}'
                        )
                        
                        # Reemplazar en delta con URL
                        op['insert']['image'] = f"/document_bp/api/image/{filename}"
                        
                        logger.info(f"Imagen subida: {filename}")
                        
                    except Exception as e:
                        logger.error(f"Error subiendo imagen: {e}")
    
    return delta

def optimize_image(image_bytes, format_ext, max_size=(1920, 1920), quality=85):
    """Optimizar imagen para reducir tamaño"""
    try:
        # Abrir imagen
        image = Image.open(BytesIO(image_bytes))
        
        # Convertir a RGB si es necesario
        if image.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', image.size, (255, 255, 255))
            background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
            image = background
        
        # Redimensionar si es muy grande
        if image.size[0] > max_size[0] or image.size[1] > max_size[1]:
            image.thumbnail(max_size, Image.Resampling.LANCZOS)
        
        # Guardar optimizada
        output_buffer = BytesIO()
        save_format = 'JPEG' if format_ext.lower() in ['jpg', 'jpeg'] else format_ext.upper()
        
        save_kwargs = {'format': save_format}
        if save_format == 'JPEG':
            save_kwargs['quality'] = quality
            save_kwargs['optimize'] = True
        
        image.save(output_buffer, **save_kwargs)
        
        return output_buffer.getvalue()
    
    except Exception as e:
        logger.error(f"Error optimizando imagen: {e}")
        return image_bytes

def save_to_minio_compressed(delta, html):
    """Guardar contenido comprimido en Minio"""
    content = {
        'delta': delta,
        'html': html,
        'timestamp': datetime.utcnow().isoformat(),
        'version': '1.0'
    }
    
    json_content = json.dumps(content, ensure_ascii=False)
    compressed = gzip.compress(json_content.encode('utf-8'))
    
    filename = f"doc_{uuid.uuid4()}.json.gz"
    
    minio_client.put_object(
        bucket_name='documents',
        object_name=filename,
        data=BytesIO(compressed),
        length=len(compressed),
        content_type='application/gzip'
    )
    
    return filename

def load_from_minio_compressed(filename):
    """Cargar contenido comprimido desde Minio"""
    try:
        response = minio_client.get_object('documents', filename)
        compressed_data = response.read()
        
        # Descomprimir
        json_content = gzip.decompress(compressed_data).decode('utf-8')
        content = json.loads(json_content)
        
        return content.get('delta', {}), content.get('html', '')
        
    except Exception as e:
        logger.error(f"Error cargando desde Minio: {e}")
        return None, None

def process_docx_upload(file_path):
    """Procesar archivo DOCX subido y convertir a formato Quill, incluyendo imágenes"""
    try:
        with open(file_path, 'rb') as docx_file:
            result = mammoth.convert_to_html(
                docx_file,
                convert_image=mammoth.images.inline(convert_image_to_base64)
            )

            html_content = result.value

            # Warnings de Mammoth
            if result.messages:
                for message in result.messages:
                    logger.warning(f"Mammoth warning: {message}")

        processed_html = process_html_for_quill(html_content)
        delta = html_to_basic_delta(processed_html)

        return delta, processed_html, None

    except Exception as e:
        logger.error(f"Error procesando DOCX: {e}")
        return None, None, str(e)

def convert_image_to_base64(image):
    """
    Convertir imágenes DOCX a base64 para incrustar en HTML.
    Mammoth pasa un objeto Image con método open() y content_type.
    """
    with image.open() as image_bytes:
        import base64
        encoded = base64.b64encode(image_bytes.read()).decode('utf-8')
        return {"src": f"data:{image.content_type};base64,{encoded}"}

def process_html_for_quill(html_content):
    """Procesar HTML para que sea compatible con Quill"""
    if not html_content:
        return ""

    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.find_all(True):
        style = element.get('style', '')
        if 'text-align: center' in style:
            element['class'] = element.get('class', []) + ['ql-align-center']
        elif 'text-align: right' in style:
            element['class'] = element.get('class', []) + ['ql-align-right']
        elif 'text-align: justify' in style:
            element['class'] = element.get('class', []) + ['ql-align-justify']

        if element.get('style'):
            del element['style']

    return str(soup)

def html_to_basic_delta(html_content):
    """Convertir HTML procesado a Delta para Quill, incluyendo imágenes"""
    if not html_content:
        return {"ops": [{"insert": "\n"}]}

    soup = BeautifulSoup(html_content, 'html.parser')
    ops = []

    def process_element(node):
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                ops.append({"insert": text})
            return
        if not isinstance(node, Tag):
            return

        if node.name == 'p':
            text = node.get_text()
            attributes = {}
            classes = node.get('class', [])
            if 'ql-align-center' in classes:
                attributes['align'] = 'center'
            elif 'ql-align-right' in classes:
                attributes['align'] = 'right'
            elif 'ql-align-justify' in classes:
                attributes['align'] = 'justify'

            if text.strip():
                ops.append({"insert": text})
                if attributes:
                    ops.append({"insert": "\n", "attributes": attributes})
                else:
                    ops.append({"insert": "\n"})

        elif node.name in ['h1','h2','h3','h4','h5','h6']:
            level = int(node.name[1])
            text = node.get_text()
            if text.strip():
                ops.append({"insert": text})
                ops.append({"insert": "\n", "attributes": {"header": level}})

        elif node.name in ['strong','b']:
            text = node.get_text()
            if text.strip():
                ops.append({"insert": text, "attributes": {"bold": True}})
        elif node.name in ['em','i']:
            text = node.get_text()
            if text.strip():
                ops.append({"insert": text, "attributes": {"italic": True}})
        elif node.name == 'u':
            text = node.get_text()
            if text.strip():
                ops.append({"insert": text, "attributes": {"underline": True}})

        elif node.name == 'img':
            src = node.get('src')
            if src:
                ops.append({"insert": {"image": src}})

        for child in node.children:
            process_element(child)

    for el in soup.contents:
        process_element(el)

    if not ops or ops[-1].get('insert') != '\n':
        ops.append({"insert": "\n"})

    return {"ops": ops}

def create_version_backup(document):
    """Crear respaldo de versión del documento"""
    
    try:
        # Contar versiones existentes
        version_count = DocumentVersion.query.filter_by(document_id=document.id).count()
        max_versions = current_app.config.get('KEEP_VERSIONS', 10)
        
        # Eliminar versiones antiguas si exceden el límite
        if version_count >= max_versions:
            oldest_versions = DocumentVersion.query.filter_by(document_id=document.id)\
                .order_by(DocumentVersion.created_at)\
                .limit(version_count - max_versions + 1).all()
            
            for old_version in oldest_versions:
                if old_version.minio_path:
                    try:
                        minio_client.remove_object('documents', old_version.minio_path)
                    except:
                        pass
                db.session.delete(old_version)
        
        # Crear nueva versión
        version = DocumentVersion(
            document_id=document.id,
            version_number=version_count + 1,
            content_delta=document.content_delta,
            content_html=document.content_html,
            minio_path=document.minio_path,
            size_bytes=document.size_bytes
        )
        
        db.session.add(version)
        db.session.commit()
        
        logger.info(f"Versión de respaldo creada para documento {document.id}")
        return version
        
    except Exception as e:
        logger.error(f"Error creando versión de respaldo: {e}")
        return None

# Funciones de exportación
def clean_html_for_export(html):
    """Limpiar HTML para exportación"""
    if not html:
        return ""
    
    soup = BeautifulSoup(html, 'html.parser')
    
    # Convertir imágenes de Minio a URLs completas
    for img in soup.find_all('img'):
        src = img.get('src', '')
        if src.startswith('/document_bp/api/image/'):
            filename = src.replace('/document_bp/api/image/', '')
            # Generar URL temporal de Minio
            try:
                img_url = minio_client.presigned_get_object(
                    'images', filename, expires=timedelta(hours=1)
                )
                img['src'] = img_url
            except:
                # Si falla, mantener URL relativa
                pass
    
    return str(soup)

def export_to_pdf(html_content, title="Documento"):
    """Exportar HTML a PDF usando WeasyPrint"""
    try:
        # CSS mejorado para el PDF
        css_content = """
        @page {
            margin: 2.5cm;
            size: A4;
            @top-center {
                content: "{title}";
                font-family: Arial, sans-serif;
                font-size: 10px;
                color: #666;
            }
            @bottom-center {
                content: "Página " counter(page) " de " counter(pages);
                font-family: Arial, sans-serif;
                font-size: 10px;
                color: #666;
            }
        }
        
        body {
            font-family: 'Segoe UI', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            font-size: 12px;
        }
        
        h1, h2, h3, h4, h5, h6 { 
            color: #2c3e50;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }
        
        h1 { font-size: 24px; }
        h2 { font-size: 20px; }
        h3 { font-size: 16px; }
        h4 { font-size: 14px; }
        h5 { font-size: 12px; }
        h6 { font-size: 11px; }
        
        p { margin-bottom: 1em; }
        
        .ql-align-center { text-align: center; }
        .ql-align-right { text-align: right; }
        .ql-align-justify { text-align: justify; }
        
        blockquote {
            border-left: 4px solid #3498db;
            padding-left: 20px;
            margin-left: 0;
            font-style: italic;
            background-color: #f8f9fa;
            padding: 15px 20px;
        }
        
        ul, ol {
            margin-bottom: 1em;
            padding-left: 2em;
        }
        
        li { margin-bottom: 0.3em; }
        
        img {
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1em auto;
        }
        
        table {
            border-collapse: collapse;
            width: 100%;
            margin-bottom: 1em;
        }
        
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        
        th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
        """.format(title=title)
        
        html_doc = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{title}</title>
        </head>
        <body>
            <h1 style="text-align: center; border-bottom: 2px solid #3498db; padding-bottom: 10px;">{title}</h1>
            {clean_html_for_export(html_content)}
        </body>
        </html>
        """
        
        pdf_buffer = BytesIO()
        HTML(string=html_doc).write_pdf(
            pdf_buffer, 
            stylesheets=[CSS(string=css_content)],
            presentational_hints=True
        )
        pdf_buffer.seek(0)
        
        return pdf_buffer
        
    except Exception as e:
        logger.error(f"Error exportando a PDF: {e}")
        return None

def export_to_docx(html_content, title="Documento"):
    """Exportar HTML a DOCX mejorado"""
    try:
        doc = DocxDocument()
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Agregar título
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar fecha
        date_paragraph = doc.add_paragraph(f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar separador
        doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Línea en blanco
        
        # Procesar HTML
        soup = BeautifulSoup(clean_html_for_export(html_content), 'html.parser')
        
        # Función recursiva mejorada para procesar elementos
        def process_element(element, parent_paragraph=None):
            if element.name == 'p':
                p = doc.add_paragraph()
                
                # Verificar alineación
                classes = element.get('class', [])
                if 'ql-align-center' in classes:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'ql-align-right' in classes:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif 'ql-align-justify' in classes:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                process_inline_elements(element, p)
                
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                heading = doc.add_heading(element.get_text().strip(), level)
                
            elif element.name == 'blockquote':
                p = doc.add_paragraph()
                p.style = doc.styles['Quote']
                process_inline_elements(element, p)
                
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text().strip(), style='List Bullet')
                    
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text().strip(), style='List Number')
            
            elif element.name == 'br':
                if parent_paragraph:
                    parent_paragraph.add_run().add_break()
                else:
                    doc.add_paragraph()
        
        def process_inline_elements(element, paragraph):
            for child in element.children:
                if hasattr(child, 'name'):
                    if child.name in ['strong', 'b']:
                        run = paragraph.add_run(child.get_text())
                        run.bold = True
                    elif child.name in ['em', 'i']:
                        run = paragraph.add_run(child.get_text())
                        run.italic = True
                    elif child.name == 'u':
                        run = paragraph.add_run(child.get_text())
                        run.underline = True
                    elif child.name == 'img':
                        # Manejar imagen como buffer
                        src = child.get('src')
                        if src:
                            try:
                                if src.startswith('http'):  # URL
                                    import requests
                                    resp = requests.get(src)
                                    image_bytes = BytesIO(resp.content)
                                    paragraph.add_run().add_picture(image_bytes, width=Inches(4))
                                else:  # Minio u otro path local
                                    paragraph.add_run().add_picture(src, width=Inches(4))
                            except Exception as e:
                                logger.warning(f"No se pudo agregar imagen: {e}")
                    elif child.name == 'br':
                        paragraph.add_run().add_break()
                    else:
                        process_inline_elements(child, paragraph)
                else:
                    text = str(child).strip()
                    if text:
                        paragraph.add_run(text)
        
        # Procesar todo el contenido
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'blockquote', 'ul', 'ol']):
            process_element(element)
        
        # Agregar pie de página
        doc.add_paragraph()
        footer = doc.add_paragraph("_" * 50)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Guardar en buffer
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        return docx_buffer
        
    except Exception as e:
        logger.error(f"Error exportando a DOCX: {e}")
        return None

# Funciones de Redis para cache y sesiones
def cache_document(doc_id, data, expire_time=300):
    """Cachear documento en Redis"""
    try:
        cache_key = f"document:{doc_id}"
        redis_client.setex(cache_key, expire_time, json.dumps(data))
        return True
    except Exception as e:
        logger.error(f"Error cacheando documento: {e}")
        return False

def get_cached_document(doc_id):
    """Obtener documento del cache"""
    try:
        cache_key = f"document:{doc_id}"
        cached_data = redis_client.get(cache_key)
        if cached_data:
            return json.loads(cached_data)
        return None
    except Exception as e:
        logger.error(f"Error obteniendo documento del cache: {e}")
        return None

def invalidate_document_cache(doc_id):
    """Invalidar cache de documento"""
    try:
        cache_key = f"document:{doc_id}"
        redis_client.delete(cache_key)
        return True
    except Exception as e:
        logger.error(f"Error invalidando cache: {e}")
        return False

def set_autosave_lock(doc_id, user_email, lock_time=30):
    """Establecer bloqueo para auto-guardado"""
    try:
        lock_key = f"autosave_lock:{doc_id}"
        lock_data = {
            'user_email': user_email,
            'timestamp': datetime.utcnow().isoformat()
        }
        redis_client.setex(lock_key, lock_time, json.dumps(lock_data))
        return True
    except Exception as e:
        logger.error(f"Error estableciendo bloqueo de auto-guardado: {e}")
        return False

def get_autosave_lock(doc_id):
    """Obtener información del bloqueo de auto-guardado"""
    try:
        lock_key = f"autosave_lock:{doc_id}"
        lock_data = redis_client.get(lock_key)
        if lock_data:
            return json.loads(lock_data)
        return None
    except Exception as e:
        logger.error(f"Error obteniendo bloqueo de auto-guardado: {e}")
        return None

# Funciones de email
def send_share_notification_email(recipient_email, document_title, shared_by_email, share_url, message=None):
    """Enviar notificación de documento compartido"""
    try:
        subject = f"Documento compartido: {document_title}"
        
        # Crear contenido HTML del email
        html_body = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
                .header {{ background: #f8f9fa; padding: 20px; border-radius: 8px; text-align: center; }}
                .content {{ padding: 20px; }}
                .button {{ 
                    display: inline-block; 
                    padding: 12px 24px; 
                    background: #007bff; 
                    color: white; 
                    text-decoration: none; 
                    border-radius: 5px; 
                    margin: 20px 0;
                }}
                .footer {{ font-size: 12px; color: #666; margin-top: 30px; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h2>📄 Documento Compartido</h2>
                </div>
                
                <div class="content">
                    <p>Hola,</p>
                    
                    <p><strong>{shared_by_email}</strong> ha compartido contigo el documento <strong>"{document_title}"</strong>.</p>
                    
                    {f'<p><em>Mensaje:</em> {message}</p>' if message else ''}
                    
                    <p>Puedes acceder al documento haciendo clic en el siguiente enlace:</p>
                    
                    <div style="text-align: center;">
                        <a href="{share_url}" class="button">Ver Documento</a>
                    </div>
                    
                    <p><small>Este enlace es válido por 7 días. Si tienes problemas para acceder, contacta a {shared_by_email}.</small></p>
                </div>
                
                <div class="footer">
                    <p>Este es un mensaje automático del Editor de Documentos. Por favor, no respondas a este correo.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        # Crear mensaje
        msg = Message(
            subject=subject,
            recipients=[recipient_email],
            html=html_body
        )
        
        # Enviar email
        mail.send(msg)
        logger.info(f"Email de compartir enviado a {recipient_email}")
        return True
        
    except Exception as e:
        logger.error(f"Error enviando email de compartir: {e}")
        return False

def format_bytes(bytes_value):
    """Formatear bytes en formato legible"""
    if bytes_value == 0:
        return "0 bytes"
    
    sizes = ['bytes', 'KB', 'MB', 'GB', 'TB']
    i = 0
    while bytes_value >= 1024 and i < len(sizes) - 1:
        bytes_value /= 1024.0
        i += 1
    
    return f"{bytes_value:.2f} {sizes[i]}"

def generate_share_url(request, share_token):
    """Generar URL completa para compartir documento"""
    base_url = f"{request.scheme}://{request.host}"
    return f"{base_url}/shared/{share_token}"